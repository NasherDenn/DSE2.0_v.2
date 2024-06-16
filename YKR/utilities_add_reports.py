# -*- coding: utf-8 -*-

from docx import Document
import logging
import traceback
import re
from YKR.props import *
from collections import Counter
import sys

# получаем имя машины с которой был осуществлён вход в программу
uname = os.environ.get('USERNAME')
# инициализируем logger
logger = logging.getLogger()
logger_with_user = logging.LoggerAdapter(logger, {'user': uname})


# получение пути и названий репортов и БД для дальнейшей работы
def get_name_dir(dir_files: str, name_dir_files: list) -> list:
    # переменная-список для дальнейшего преобразования списка списков в список строк выбранных для загрузки файлов docx
    name_dir_docx = []
    for i in name_dir_files:
        name_dir_docx.append(f'{dir_files}{i}')
    return name_dir_docx


# проверяем наличие всех БД (с 2019 по 2026 года) во всех вариациях в папке "DB"
def db_in_folder():
    path_db = f'{os.path.abspath(os.getcwd())}\\DB\\'
    name_dir_db = []
    for (dirpath, dirnames, filenames) in os.walk(path_db):
        name_dir_db.extend(filenames)
    # полученный список БД
    list_db_for_check = get_name_dir(path_db, name_dir_db)
    # список баз данных, которых нет в папке "DB"
    return list(set(list_db) - set(list_db_for_check))


def get_dirty_data_report(path_to_report: str, report_number: str) -> dict:
    doc = Document(path_to_report)
    # переменная со всеми таблицами в репорте
    all_tables = doc.tables
    # создаем пустой словарь под неочищенные данные таблиц
    dirty_data_tables = {i: None for i in range(0, len(all_tables))}
    for i, table in enumerate(all_tables):
        # создаем список строк для таблицы `i` (пока пустые)
        dirty_data_tables[i] = [[] for _ in range(0, len(table.rows))]
        # проходимся по строкам таблицы `i`
        for j, row in enumerate(table.rows):
            try:
                for cell in row.cells:
                    dirty_data_tables[i][j].append(cell.text)
            except IndexError:
                logger_with_user.error(f'Ошибка в {report_number}:\n {traceback.format_exc()}')
                break
    return dirty_data_tables


# выбор только файлов (репортов) в названиях которых есть "04-YKR"
def change_only_ykr_reports(name_dir_docx: list) -> list:
    list_reports_for_work = []
    not_loading_welding = []
    not_loading_shaer_wave = []
    for index, i in enumerate(name_dir_docx):
        if 'WELDING' in i.upper():
            not_loading_welding.append(index)
        if 'SHAER WAVE' in i.upper():
            not_loading_shaer_wave.append(index)
        if '04-YKR' in i.upper():
            list_reports_for_work.append(i)

    return list_reports_for_work, not_loading_welding, not_loading_shaer_wave


# получаем номер репорта, номер work order и дату
def number_report_wo_date(path_to_report: str) -> dict:
    try:
        doc = Document(path_to_report)
    except:
        logger_with_user.warning(f'Не могу обработать верхний колонтитул в отчёте {path_to_report}!!!!\n'
                                 f'{traceback.format_exc()}')
        sys.exit()
    # получаем неочищенные данные из первого верхнего колонтитула
    head_paragraph = doc.sections[0].header.tables
    # активатор, если для первого листа установлен отдельный колонтитул
    must_check_head_paragraph = False
    # если колонтитул для первого листа отдельный
    if doc.sections[0].different_first_page_header_footer:
        must_check_head_paragraph = True
    # создаем пустой словарь под данные верхнего колонтитула
    data_header = {i: None for i in range(0, len(head_paragraph))}
    for i, table in enumerate(head_paragraph):
        # создаем список строк для таблицы `i` (пока пустые)
        data_header[i] = [[] for _ in range(0, len(table.rows))]
        # проходимся по строкам таблицы `i`
        for j, row in enumerate(table.rows):
            try:
                for cell in row.cells:
                    data_header[i][j].append(cell.text)
            except:
                logger_with_user.warning(f'Не могу обработать верхний колонтитул в отчёте {path_to_report}!!!!\n'
                                         f'{traceback.format_exc()}')
                sys.exit()
    return get_number_report_wo_date(data_header, path_to_report), must_check_head_paragraph


# получение фактического номера репорта, номера work order и даты
def get_number_report_wo_date(data_header: dict, path_to_report: str) -> dict:
    # словарь номера репорта, даты репорта, номера Work Order
    rep_number = {}
    # перебираем полученные данные из первого верхнего колонтитула
    # проходим по нему как по таблице
    for index_row, i in enumerate(data_header[0]):
        for index_column, ii in enumerate(i):
            try:
                # если слова "Report No:" и фактический номер репорта в одной ячейке
                if 'Report' in ii and any(map(str.isdigit, data_header[0][index_row][index_column])):
                    # индекс начала номера репорта в той же ячейке
                    index_start_number_report = re.search("\\d", ii).start()
                    rep_number['report_number'] = data_header[0][index_row][index_column][index_start_number_report:]
                # иначе если слова "Report No:" и фактический номер репорта в разных ячейках
                elif 'Report' in ii and any(map(str.isdigit, data_header[0][index_row][index_column + 1])):
                    # индекс начала номера репорта в соседней ячейке
                    index_start_number_report = re.search("\\d", data_header[0][index_row][index_column + 1]).start()
                    rep_number['report_number'] = data_header[0][index_row][index_column + 1][index_start_number_report:]
                # если слово "Date" и фактическая дата репорта в одной ячейке
                if 'Date' in ii and any(map(str.isdigit, data_header[0][index_row][index_column])):
                    # индекс начала даты репорта в той же ячейке
                    index_start_number_report = re.search("\\d", ii).start()
                    rep_number['report_date'] = data_header[0][index_row][index_column][index_start_number_report:]
                    # приводим номер репорта к стандарту
                    rep_number['report_date'] = format_date_to_normal_form(rep_number['report_date'])
                # иначе если слово "Date" и фактическая дата репорта в разных ячейках
                elif 'Date' in ii and any(map(str.isdigit, data_header[0][index_row][index_column + 1])):
                    # индекс начала даты репорта в соседней ячейке
                    index_start_number_report = re.search("\\d", data_header[0][index_row][index_column + 1]).start()
                    rep_number['report_date'] = data_header[0][index_row][index_column + 1][index_start_number_report:]
                    # приводим номер репорта к стандарту
                    rep_number['report_date'] = format_date_to_normal_form(rep_number['report_date'])
                # если слово "order" и номер work order в одной ячейке
                if 'order' in ii and any(map(str.isdigit, data_header[0][index_row][index_column])):
                    # индекс начала номера work order в той же ячейке
                    index_start_number_report = re.search("\\d", ii).start()
                    rep_number['work_order'] = data_header[0][index_row][index_column][index_start_number_report:]
                elif 'order' in ii and 'N/A' in ii.upper():
                    rep_number['work_order'] = 'n/a'
                elif 'REQU' in ii.upper():
                    rep_number['work_order'] = 'NCOC Request'
                # иначе если слова "order" и номер work order в разных ячейках
                elif 'order' in ii and any(map(str.isdigit, data_header[0][index_row][index_column + 1])):
                    # индекс начала номера work order в соседней ячейке
                    index_start_number_report = re.search("\\d", data_header[0][index_row][index_column + 1]).start()
                    rep_number['work_order'] = data_header[0][index_row][index_column + 1][index_start_number_report:]
                # иначе если слова "order" и номер work order в разных ячейках и нет цифр, значит номер work order - NCOC Request
                elif 'order' in ii and not any(map(str.isdigit, data_header[0][index_row][index_column + 1])):
                    rep_number['work_order'] = data_header[0][index_row][index_column + 1]
            except:
                logger_with_user.warning(f'Что-то не так в верхнем колонтитуле файла {path_to_report}\n'
                                         f'{traceback.format_exc()}')
                sys.exit()
    # возвращаем не очищенные значения номера репорта, номера work order и даты
    return rep_number


# очистка номера репорта, даты репорта, номера Work Order от лишних, повторяющихся символов
def clear_data_rep_number(data: dict, must_check: bool) -> dict:
    # удаление любых пробельных символов в номере репорта
    data['report_number'] = re.sub('\\s+', '', data['report_number'])
    # замена повторяющегося символа "-" на единичный в номере репорта
    data['report_number'] = re.sub('-+', '-', data['report_number'])
    # замена любых пробельных символов в дате репорта на "."
    data['report_date'] = re.sub('\\s+', '.', data['report_date'])
    # замена повторяющегося символа "." на единичный в дате репорта
    data['report_date'] = re.sub('\\.+', '.', data['report_date'])
    # замена повторяющегося символа "-" на единичный в дате репорта
    data['report_date'] = re.sub('-+', '-', data['report_date'])
    # удаление любых пробельных символов в work order
    data['work_order'] = re.sub('\\s+', '', data['work_order'])
    # если в номере репорта была Revision
    if 'Rev' in data['report_number'] or 'rev' in data['report_number'] or 'REV' in data['report_number']:
        # то добавляем "Rev." через знак "_"
        index_rev = data['report_number'].find('ev')
        data['report_number'] = '_'.join([data['report_number'][:index_rev - 1], data['report_number'][index_rev - 1:]])
    # if must_check:
    #     num_rep = data['report_number']
    #     logger_with_user.info(f'Проверь верхние колонтитулы всех листов отчёта {num_rep}, т.к. установлена опция "Особый колонтитул для первой '
    #                           f'страницы". Номера отчётов на страницах могут отличаться!')
    return data


# вытягиваем из номера репорта локацию (ON, OF, OS), метод контроля (UTT, PAUT), год контроля (18, 19, 20, 21, 22, 23, 24, 25, 26)
# и формирование имени БД для дальнейшей записи
def reports_db(name_report: str, break_break: bool) -> tuple:
    location = ['-ON-', '-on-', '-OF-', '-of-', '-OFF-', '-off-', '-OS-', '-os-']
    method = ['-UT-', '-ut-', '-UTT-', '-utt-', '-PAUT-', '-paut-']
    years = ['-18-', '-19-', '-20-', '-21-', '-22-', '-23-', '-24-', '-25-', '-26-']
    name_for_reports_db = ''
    if break_break:
        for i in location:
            if i in name_report:
                name_for_reports_db = f'reports_db_{i[1:-1]}_'
                break
        if name_for_reports_db == '':
            name_for_reports_db = 'reports_db_ON_'
        # активатор, если не нашли метод контроля в номере репорта
        find = False
        for i in method:
            if i in name_report:
                name_for_reports_db = f'{name_for_reports_db}{i[1:-1].upper()}_'
                if '_UT_' in name_for_reports_db:
                    name_for_reports_db = name_for_reports_db.replace('_UT_', '_UTT_')
                if '_OFF_' in name_for_reports_db:
                    name_for_reports_db = name_for_reports_db.replace('_OFF_', '_OF_')
                find = True
        # если не нашли метод контроля, то переходим к следующему репорту
        if not find:
            logger_with_user.error(f'Не могу определить метод контроля! Проверь корректность записи номера репорта {name_report}!1')
            break_break = False
        if break_break:
            find = False
            for i in years:
                if i in name_report:
                    name_for_reports_db = f'{name_for_reports_db}{i[1:-1]}.sqlite'
                    find = True
            if not find:
                logger_with_user.error(f'Не могу определить год контроля! Проверь корректность записи номера репорта {name_report}!2')
                break_break = False
    return name_for_reports_db, break_break


# Получаем номера словарей (таблиц) в которых есть ключевое слово "Nominal thickness".
# Если при переборе ячеек таблицы найден "Nominal thickness" и "Project", то делаем запись в logger и исключаем эту таблицу
# из дальнейшей обработки.
# При этом проверяем, есть ли в ней реальная таблица с данными на основании нахождения в ней одного значения из списка (North, South, East,
# Row, West, Extrados, Intrados, Central и т.д.)
# На выходе получаем номер таблицы для дальнейшей обработки и список номеров таблиц с реальными данными (номера для записи в БД).
def first_clear_table_nominal_thickness(first_dirty_table: dict or list, number_dirty_table: int, report_number: str, method: str) -> int or str:
    # перебираем строки в словаре (таблице)
    # активатор наличия "Nominal thickness"
    nominal_is = False
    # активатор наличия "Project"
    project_is = False
    # активатор наличия одного значения из списка (North, South, East, Row, West, Extrados, Intrados, Central и т.д.)
    name_column_is = False
    utt_one_of_name_column = ['NORTH', 'SOUTH', 'WEST', 'EAST', 'SECT', 'EXTRADOS', 'INTRADOS', 'ROW', 'COLUMN', 'SPOT', 'ISOM', 'P&ID', 'S/NO',
                              'O\'CLOCK',
                              'CENTER', 'LOC', 'CONTR']
    paut_one_of_name_column = ['START X', 'END X', 'AVERAGE']

    for index_row, row in enumerate(first_dirty_table):
        index_row_nominal_is = -1000
        # перебираем колонки в строке
        for column in row:
            if 'NOM' in column.upper():
                nominal_is = True
                index_row_nominal_is = index_row
            if 'PROJ' in column.upper():
                project_is = True
            if method == 'utt':
                for word in utt_one_of_name_column:
                    if word in column.upper():
                        index_row_name_column_is = index_row
                        if index_row_nominal_is == index_row_name_column_is:
                            name_column_is = True
            if method == 'paut':
                for word in paut_one_of_name_column:
                    if word in column.upper():
                        index_row_name_column_is = index_row
                        if index_row_nominal_is == index_row_name_column_is:
                            name_column_is = True
    # если в одной таблице и "Nominal thickness", и "Project", и одно значение названия столбца из списка выше
    if nominal_is and project_is and name_column_is:
        logger_with_user.warning(f'Проверь репорт {report_number}! Первая таблица с рабочей информацией не отделена от таблиц(ы) с данными')
        return str(number_dirty_table)
    # если в таблице только "Nominal thickness"
    if nominal_is and not project_is:
        return number_dirty_table


# удаляем строки в которых есть слова "result", "details", "notes"
def delete_first_string(second_dirty_table: list) -> list:
    # список номеров строк для удаления
    index_delete_string = []
    for i, row in enumerate(second_dirty_table):
        for column in row:
            if 'RES' in column.upper() or 'DET' in column.upper() or 'NOTE' in column.upper() or 'ANEX' in column.upper() \
                    or 'ANNEX' in column.upper():
                index_delete_string.append(i)
    if index_delete_string:
        # удаляем повторяющиеся номера
        index_delete_string = list(set(index_delete_string))
        # сортируем номера по убыванию
        index_delete_string.sort(reverse=True)
        for index in index_delete_string:
            second_dirty_table.pop(index)
    return second_dirty_table


# Проверяем таблицы (словарь таблиц), что бы в каждой строке было одинаковое количество ячеек.
# Если нет, то в таблице есть сдвиги полей, т.е. таблица геометрически не ровная.
def check_len_row(table_uncheck_len_row: dict, report_number: str) -> dict:
    # список номеров таблиц для удаления из словаря
    list_number_table_unequal_row = []
    for i in table_uncheck_len_row.keys():
        # список длин строк в таблице
        list_len_row = []
        for row in table_uncheck_len_row[i]:
            list_len_row.append(len(row))
        if len(set(list_len_row)) != 1 or not set(list_len_row):
            list_number_table_unequal_row.append(i)
    if list_number_table_unequal_row:
        for i in list_number_table_unequal_row:
            logger_with_user.warning(f'"Не ровная" таблица. Проверь таблицу {i} в репорте {report_number}')
            table_uncheck_len_row.pop(i)
    return table_uncheck_len_row


# определяем чем является таблица: "сетка" или обыкновенная
# на выходе получаем словарь {"mesh": ["номер таблиц"], "ord": ["номер таблиц"]}
# "mesh" - сетка, "ord" - обыкновенная
def which_table(data_table_equal_row: dict) -> list:
    # список таблиц, которые являются "сеткой"
    this_is_mesh = []
    # список таблиц, которые являются обыкновенными
    for i in data_table_equal_row.keys():
        check = 0
        # если в таблице больше 6 строк (в сетке много строк)
        if len(data_table_equal_row[i]) > 5:
            for step in range(4):
                if len(set(data_table_equal_row[i][step])) == 2:
                    check += 1
        # если хотя бы 3 строки (защита, от пропуска на предыдущем этапе отсеивания "не ровных" таблиц) имеют две ячейки, то - "сетка"
        if check > 2:
            this_is_mesh.append(i)
    return this_is_mesh


# преобразуем таблицы с "сеткой" - переносим первые четыре строки в названия столбцов и их значения
# на входе словарь после преобразования ("не ровная" таблица) и список таблиц "сетка"
# на выходе новый словарь преобразованных таблиц: первая строка - название столбцов, остальные - данные
def converted_mesh(data_table_equal_row: dict, mesh_table: list, number_report: str) -> dict:
    for index_table in mesh_table:
        # словарь {"название столбцов": "значение"} для дальнейшего преобразования
        name_value_converted = {}
        # перебираем первые 5 строк в таблице
        for step in range(6):
            # если длина уникальных значений в строке == 2
            if len(set(data_table_equal_row[index_table][step])) == 2:
                # то перебираем их в виде списка
                name_value = list(set(data_table_equal_row[index_table][step]))
                for future_column in name_value:
                    # и сравниваем со списком возможных названий столбцов
                    if 'LIN' in future_column.upper() or 'TAG' in future_column.upper() or 'CONT' in future_column.upper() \
                            or 'DRAW' in future_column.upper() or 'ISOM' in future_column.upper():
                        name_column = 'Line'
                    elif 'DIA' in future_column.upper() or 'INCH' in future_column.upper():
                        name_column = 'Diameter'
                    elif 'NOM' in future_column.upper():
                        name_column = 'Nominal_thickness'
                    elif 'ITEM' in future_column.upper() or 'DESC' in future_column.upper():
                        name_column = 'Item_description'
                    else:
                        value = future_column
                try:
                    name_value_converted[name_column] = value
                except UnboundLocalError:
                    logger_with_user.error(f'Не могу записать таблицы в репорте {number_report}\n'
                                           f'{traceback.format_exc()}')
                    sys.exit()
        # проверка на то, что мы получили все четыре строки (Line, Diameter, Nominal_thickness, Item_description)
        if len(name_value_converted) != 4:
            logger_with_user.error(f'В таблице {index_table} репорта {number_report} не могу найти Line или Diameter, '
                                   f'или Nominal_thickness, или Item_description')
            continue
        # проверяем находится ли в одной ячейке номер линии и чертежа
        if re.findall(r'[AАBВCСDHНMМ]'
                      r'\d{1,2}'
                      r'-{1,2}?\s?'
                      r'\d{3,4}'
                      r'-?\s?'
                      r'\D{2}'
                      r'-?\s?'
                      r'\d{3}', name_value_converted['Line']) and re.findall(r'[KК][EЕ]01[-.].+|[TТ]R01[-.].+', name_value_converted['Line']):
            # если да, то рассоединяем их
            line_drawing = disconnect_line_drawing(name_value_converted['Line'])
            # заменяем 'Line' и дополняем 'Drawing' в name_value_converted
            name_value_converted['Line'] = line_drawing['Line']
            name_value_converted['Drawing'] = line_drawing['Drawing']

        # на данном этапе получили словарь name_value_converted типа:
        # {'Line': 'A1-2002-RO-007-20-A11-HC', 'Diameter': '8”', 'Nominal_thickness': '8.18 mm', 'Item_description': 'TML-001 (Elbow)',
        #  'Drawing': 'KE01-A1-200-PO-P-DI-2086-001'}
        # удаляем из каждой таблицы (data_table_equal_row[index_table]) первые 4 строки с номером линии, чертежа, диаметром, объектом
        data_table_equal_row[index_table] = data_table_equal_row[index_table][4:]
        # теперь в таблице data_table_equal_row[index_table] первый список (строка) - название колонок, остальные - данные
        # соединяем ключи name_value_converted с первым списком data_table_equal_row[index_table],
        # значения name_value_converted с каждым последующим списком в data_table_equal_row[index_table]
        # 'Line' - всегда первое место, затем 'Drawing', 'Item_description', 'Diameter', 'Nominal_thickness'

        # дополняем первую строку (список) в каждой таблице новыми названиями столбцов
        if len(name_value_converted) == 5:
            data_table_equal_row[index_table][0].insert(0, 'Line')
            data_table_equal_row[index_table][0].insert(1, 'Drawing')
            data_table_equal_row[index_table][0].insert(2, 'Item_description')
            data_table_equal_row[index_table][0].insert(3, 'Diameter')
            data_table_equal_row[index_table][0].insert(4, 'Nominal_thickness')
        if len(name_value_converted) == 4:
            data_table_equal_row[index_table][0].insert(0, 'Line')
            data_table_equal_row[index_table][0].insert(1, 'Item_description')
            data_table_equal_row[index_table][0].insert(2, 'Diameter')
            data_table_equal_row[index_table][0].insert(3, 'Nominal_thickness')

        # дополняем остальные строки с данными значениями новых столбцов
        # перебираем строки, начиная со второй в каждой таблице, и дополняем их значениями
        for row in data_table_equal_row[index_table][1:]:
            if len(name_value_converted) == 5:
                row.insert(0, name_value_converted['Line'])
                row.insert(1, name_value_converted['Drawing'])
                row.insert(2, name_value_converted['Item_description'])
                row.insert(3, name_value_converted['Diameter'])
                row.insert(4, name_value_converted['Nominal_thickness'])
            if len(name_value_converted) == 4:
                row.insert(0, name_value_converted['Line'])
                row.insert(1, name_value_converted['Item_description'])
                row.insert(2, name_value_converted['Diameter'])
                row.insert(3, name_value_converted['Nominal_thickness'])
    return data_table_equal_row


# рассоединение номера линии и номера чертежа
def disconnect_line_drawing(line_and_drawing: str) -> dict:
    # словарь рассоединённых и очищенных номера линии и номера чертежа
    clean_line_drawing = {}
    # получаем "грязный" номер чертежа в виде строки
    drawing = re.findall(r'[KК][EЕ]01[-.].+|[TТ]R01[-.].+', line_and_drawing)[0]
    # получаем "грязный" номер линии в виде строки
    line = line_and_drawing.replace(drawing, '')
    # очищаем их от лишних знаков
    clean_line_drawing['Line'] = dirt_cleaning(line)
    clean_line_drawing['Drawing'] = dirt_cleaning(drawing)
    return clean_line_drawing


# функция очистки строки от пустых символов, недопустимых знаков и т.д.
def dirt_cleaning(dirt_str: str) -> str:
    # удаление любых пробельных символов
    dirt_str = re.sub('\\s+', '', dirt_str)
    # удаление знака "/"
    dirt_str = re.sub('/', '', dirt_str)
    # # меняем "." на "-"
    # dirt_str = re.sub('\.', '-', dirt_str)
    dirt_str = dirt_str.upper()
    return dirt_str


# приводим в порядок названия столбцов (первый список) и данные (остальные строки)
def shit_in_shit_out(finish_dirty_table: dict, method: str, number_report: str) -> dict:
    # итоговый, очищенный, приведённый в порядок словарь {"номер таблицы": [[названия столбцов], [[данные], [данные]]]}
    finish_data = {}
    for index_table in finish_dirty_table.keys():
        # форматируем названия столбцов
        # по умолчанию берём первую строку как названия столбцов
        # надо проверять с какой строки идут названия столбцов (number_row_name_column) и номер строки с которой начинаются данные
        # и убирать пустые строки (картинки)
        # перебирать finish_dirty_table[index_table] построчно пока не будет найдена последняя строка содержащая 'Line' и 'Nom'
        if 'utt' in method:
            try:
                number_row_name_column = search_number_row_name_column(finish_dirty_table[index_table], method)
                number_row_data = number_row_name_column + 1
                # print(f'number_row_name_column {number_row_name_column}')
                # print(f'number_row_data {number_row_data}')
            except TypeError:
                # если не найдено ни одно ключевое слово из возможных названий столбцов
                logger_with_user.error(f'Не могу записать таблицу {index_table} репорта {number_report}. Проверь ключевые слова для поиска')
                continue
        if 'paut' in method:
            try:
                number_row_name_column = search_number_row_name_column(finish_dirty_table[index_table], method) + 1
                # print(f'number_row_name_column {number_row_name_column}')
            except TypeError:
                # если не найдено ни одно ключевое слово из возможных названий столбцов
                logger_with_user.error(f'Не могу записать таблицу {index_table} репорта {number_report}. Проверь ключевые слова для поиска \n'
                                       f'{traceback.format_exc()}')
                continue
            number_row_data = number_row_name_column + 1
        try:
            finish_name_column = cleaning_name_column(finish_dirty_table[index_table][number_row_name_column], method)
            # print(finish_name_column)
        except IndexError:
            logger_with_user.error(f'Не могу записать таблицу {index_table} репорта {number_report} \n'
                                   f'Таблица будет записана не правильно или не полностью\n'
                                   f'{traceback.format_exc()}')
        # форматируем значения данных во всех остальных строках
        try:
            finish_value_table = cleaning_value_table(finish_dirty_table[index_table][number_row_data:])
        except UnboundLocalError:
            logger_with_user.error(f'Не могу записать таблицу {index_table} репорта {number_report} \n'
                                   f'Таблица будет записана не правильно или не полностью\n'
                                   f'{traceback.format_exc()}')
        try:
            finish_data[index_table] = [finish_name_column, finish_value_table]
        except UnboundLocalError:
            logger_with_user.error(f'Не могу записать таблицу {index_table} репорта {number_report} \n'
                                   f'Таблица будет записана не правильно или не полностью\n'
                                   f'{traceback.format_exc()}')
    # print(finish_data)
    return finish_data


# поиск последнего номера строки, которая является названием столбцов перед строками с данными
def search_number_row_name_column(table: list, method: str) -> int:
    if method == 'utt':
        check_index_row = 0
        last_number_name_column = 0
        for index_row, row in enumerate(table):
            # наличие в строке слова 'Line'
            line_in_row = False
            # наличие в строке слова 'Nominal thickness'
            nominal_in_row = False
            for column in row:
                # дополнить перебор возможными словами
                if 'LINE' in column.upper() or 'ITEM' in column.upper() or 'NORTH' in column.upper() or 'TOP' in column.upper() \
                        or 'INTRADOS' in column.upper() or 'O\'CLOCK' in column.upper() or 'S/N' in column.upper() or 'START' in column.upper() \
                        or 'END' in column.upper():
                    line_in_row = True
                    # continue
                if 'NOM' in column.upper():
                    nominal_in_row = True
                    # continue
                if line_in_row and nominal_in_row:
                    last_number_name_column = index_row
                    # return last_number_name_column
                    check_index_row = index_row
            if check_index_row - index_row != 0:
                # print(f'check_index_row {check_index_row}')
                # print(f'index_row {index_row}')
                # print(f'last_number_name_column {last_number_name_column}')
                # print(last_number_name_column)
                return last_number_name_column
    if method == 'paut':
        # print('it\'s paut')
        for index_row, row in enumerate(table):
            # наличие в строке слова 'Line'
            line_in_row = False
            # наличие в строке слова 'Nominal thickness'
            nominal_in_row = False
            # print(row)
            for column in row:
                # дополнить перебор возможными словами
                if 'LIN' in column.upper() or 'CONT' in column.upper() or 'END' in column.upper() or 'AVERAGE' in column.upper():
                    line_in_row = True
                if 'NOM' in column.upper():
                    nominal_in_row = True
            if line_in_row and nominal_in_row:
                last_number_name_column = index_row
                # print(f'last_number_name_column {last_number_name_column}')
                return last_number_name_column


# приводим в порядок названия столбцов
def cleaning_name_column(list_dirty_name_column: list, method: str) -> list:
    for i, column in enumerate(list_dirty_name_column):
        # print(column)
        stop = True
        # если в названии столбца указаны "часы"
        num = False
        for element in column:
            if element.isnumeric():
                num = True
        clock = False
        if 'CLOCK' in column.upper():
            clock = True
        if clock and num:
            new_column = re.sub('\\s+', '', column)
            new_column = new_column.replace("'", '')
            list_dirty_name_column.pop(i)
            # вставляем на удалённое место новое допустимое название столбца
            list_dirty_name_column.insert(i, new_column)
        # если название столбца это только цифры
        elif re.sub('\\s', '', column).isnumeric():
            column = re.sub('\\s', '', column)
            list_dirty_name_column.pop(i)
            # вставляем на удалённое место новое допустимое название столбца
            list_dirty_name_column.insert(i, column)
        else:
            if 'utt' in method:
                # если в названии столбца присутствуют стандартные названия
                for utt_name_column_for_search in tuple_utt_name_column_for_search:
                    if stop:
                        for name in utt_name_column[utt_name_column_for_search]:
                            if name in column.upper():
                                list_dirty_name_column.pop(i)
                                # вставляем на удалённое место допустимое название столбца
                                list_dirty_name_column.insert(i, utt_name_column_for_search)
                                stop = False
                                break
                    else:
                        break
            if 'paut' in method:
                # если в названии столбца присутствуют стандартные названия
                for paut_name_column_for_search in tuple_paut_name_column_for_search:
                    if stop:
                        for name in paut_name_column[paut_name_column_for_search]:
                            if name in column.upper():
                                list_dirty_name_column.pop(i)
                                # вставляем на удалённое место допустимое название столбца
                                list_dirty_name_column.insert(i, paut_name_column_for_search)
                                stop = False
                                break
                    else:
                        break
        if stop:
            # удаляем все пробельные символы
            if re.findall('\\s|/|\.|\\(|\\)', list_dirty_name_column[i]):
                new_column = re.sub('\\s|/|\.|\\(|\\)', '', list_dirty_name_column[i])
                # new_column = re.sub('\\s', '', list_dirty_name_column[i])
                list_dirty_name_column.pop(i)
                # вставляем на удалённое место новое допустимое название столбца
                list_dirty_name_column.insert(i, new_column)
            if column == '':
                column = 'NA'
                list_dirty_name_column.pop(i)
                list_dirty_name_column.insert(i, column)
            # если первый символ цифра
            if column != '':
                if column[0].isnumeric():
                    new_column = f'_{list_dirty_name_column[i]}_'
                    list_dirty_name_column.pop(i)
                    # вставляем на удалённое место новое допустимое название столбца
                    list_dirty_name_column.insert(i, new_column)
            else:
                logger_with_user.error(f'Проверь названия столбцов {traceback.format_exc()}')
            # заменяем все возможные символ "-" ('-50', '0-50' и др.)
            if column[0] == '-':
                new_column = column.replace('-', 'minus_', 1)
                list_dirty_name_column.pop(i)
                list_dirty_name_column.insert(i, new_column)
                column = new_column
            # заменяем все возможные символ "+" ('+50')
            if column[0] == '+':
                new_column = column.replace('+', 'plus_', 1)
                list_dirty_name_column.pop(i)
                list_dirty_name_column.insert(i, new_column)
                column = new_column
            if column[-1] == '-':
                new_column = column.replace(column[-1], '_')
                list_dirty_name_column.pop(i)
                list_dirty_name_column.insert(i, new_column)
                column = new_column
            if '-' in column:
                new_column = column.replace('-', '_')
                list_dirty_name_column.pop(i)
                list_dirty_name_column.insert(i, new_column)
                column = new_column
            if ':' in column:
                new_column = column.replace(':', '_')
                list_dirty_name_column.pop(i)
                list_dirty_name_column.insert(i, new_column)
                column = new_column
            if '(' in column:
                new_column = column.replace('(', '_')
                list_dirty_name_column.pop(i)
                list_dirty_name_column.insert(i, new_column)
                column = new_column
            if ')' in column:
                new_column = column.replace(')', '_')
                list_dirty_name_column.pop(i)
                list_dirty_name_column.insert(i, new_column)
                column = new_column
            if 'mm' in column:
                new_column = column.replace('mm', '')
                list_dirty_name_column.pop(i)
                list_dirty_name_column.insert(i, new_column)
    return list_dirty_name_column


# приводим в порядок значения данных
def cleaning_value_table(list_dirty_value_table: list) -> list:
    for ii, row in enumerate(list_dirty_value_table):
        for i, column in enumerate(row):
            # new_column = re.sub(',', '.', column)
            # new_column = re.sub('\'+|”|"|’’', '', new_column)
            # new_column = re.sub('\s+', '_', new_column)
            # new_column = re.sub(':', '_', new_column)
            # new_column = re.sub('÷', '_', new_column)
            # new_column = re.sub('/', '_', new_column)
            new_column = column
            if new_column == '':
                # new_column = re.sub('', '-', column)
                new_column = re.sub('', '-', new_column)
            elif new_column == '_':
                # new_column = re.sub('_', '-', column)
                new_column = re.sub('_', '-', new_column)
            else:

                new_column = re.sub(',', '.', new_column)
                new_column = re.sub('\'+|”|"|’’', '', new_column)
                new_column = re.sub('\s+', '_', new_column)
                new_column = re.sub(':', '_', new_column)
                new_column = re.sub('÷', '_', new_column)
                new_column = re.sub('/', '_', new_column)

            row.pop(i)
            row.insert(i, new_column)
    return list_dirty_value_table


# проверяем и меняем повторяющиеся названия столбцов
def duplicate_name_column(pure_data_table: dict) -> dict:
    for number_table in pure_data_table.keys():
        all_count = Counter(pure_data_table[number_table][0])
        for count_duplicate_name_column in all_count.keys():
            if int(all_count[count_duplicate_name_column]) > 1:
                index = 1
                first_duplicate = True
                for index_column, column in enumerate(pure_data_table[number_table][0]):
                    if column == count_duplicate_name_column:
                        # если это первый дубликат, то пропускаем его (оставляем его название неизменным)
                        if first_duplicate:
                            first_duplicate = False
                            continue
                        # иначе добавляем номер в конце для корректной записи в БД
                        else:
                            pure_data_table[number_table][0][index_column] = f'{column}_{index}'
                            index += 1
    # print(f'duplicat {pure_data_table}')
    return pure_data_table


# проверяем есть ли в столбце "Line" номер чертежа, если да, то разъединяем их и дополняем новым столбцом "Drawing"
def check_drawing_in_line(pure_data_table: dict) -> dict:
    for number_table in pure_data_table.keys():
        add_column_drawing = False
        for index_row, row in enumerate(pure_data_table[number_table][1]):
            for index_column, column in enumerate(row):
                if re.findall(r'[AАBВCСDHНMМ]'
                              r'\d{1,2}'
                              r'-{1,2}?\s?'
                              r'\d{3,4}'
                              r'-?\s?'
                              r'\D{2}'
                              r'-?\s?'
                              r'\d{3}', column) and re.findall(r'[KК][EЕ]01[-.].+|[TТ]R01[-.].+', column):
                    # рассоединяем номер линии и номер чертежа
                    dict_line_drawing = disconnect_line_drawing(column)
                    dict_line_drawing['Line'] = dict_line_drawing['Line'].replace('_', '')
                    dict_line_drawing['Drawing'] = dict_line_drawing['Drawing'].replace('_', '')
                    # удаляем по индексу общую ячейку с номером линии и чертежа
                    pure_data_table[number_table][1][index_row].pop(index_column)
                    # вставляем на перове место рассоединённый номер линии, а на второе место номер чертежа
                    pure_data_table[number_table][1][index_row].insert(0, dict_line_drawing['Line'])
                    pure_data_table[number_table][1][index_row].insert(1, dict_line_drawing['Drawing'])
                    add_column_drawing = True

        # вставляем названия столбцов "Line" на первое место, "Drawing" на второе
        if add_column_drawing:
            try:
                pure_data_table[number_table][0].remove('Line')
                pure_data_table[number_table][0].insert(0, 'Line')
                pure_data_table[number_table][0].insert(1, 'Drawing')
            except:
                logger_with_user.info(f'-----------------------------------------\n'
                                      f'Критическая ошибка!!! Таблица записана не правильно!!!\n'
                                      f'{traceback.format_exc()}\n'
                                      f'-----------------------------------------')
    return pure_data_table


# список номеров таблиц, которые прошли все очистки, для дальнейшего переименования порядковых номеров таблиц для записи в БД
def take_finish_list_number_table(list_number_table: list, dict_data: dict) -> list:
    finish_list_number = list(dict_data.keys())
    for i in list_number_table:
        if type(i) == str:
            finish_list_number.insert(0, i)
    return finish_list_number


# есть ли столбец "Line" данных
def check_is_line_in_data(pure_data_table: dict) -> bool:
    line_in_data = False
    for number_table in pure_data_table.keys():
        for index_row, row in enumerate(pure_data_table[number_table][0]):
            if 'Line' in row:
                line_in_data = True
    return line_in_data


# ищем 'Line' в первой таблице и добавляем в итоговый словарь
def line_from_top_to_general_data(dirty_data_report: dict, pure_data_table: dict, report_number: str) -> dict:
    stop = True
    stop_stop = False
    # находим номер строки в которой есть ключевые слова 'Control', 'Object', 'Tag'
    if stop:
        for number_dirty_table in dirty_data_report.keys():
            if stop:
                for index, row in enumerate(dirty_data_report[number_dirty_table]):
                    if stop:
                        for cell in row:
                            if stop:
                                if 'CONTR' in cell.upper() or 'OBJEC' in cell.upper() or 'TAG' in cell.upper():
                                    stop = False
                                    # номер последнего вхождения искомого ключевого слова (следующий номер - это номер линии)
                                    index_line = dict(map(reversed, enumerate(row)))[cell]
                                    # номер вхождения номера линии в строке
                                    index_line += 1
                                    # номер строки
                                    index_row = index
                                    break
                            else:
                                break
                    else:
                        break
            else:
                break
            try:
                line_or_drawing = dirty_data_report[number_dirty_table][index_row][index_line]
            except:
                logger_with_user.error(f'В репорте {report_number} таблице {number_dirty_table} какая-то ошибка! А именно:\n'
                                       f'{traceback.format_exc()}')
    # если это линия
    try:
        if re.findall(r'[AАBВCСDHНMМ]'
                      r'\d{1,2}'
                      r'-{1,2}?\s?'
                      r'\d{3,4}'
                      r'-?\s?'
                      r'\D{2}'
                      r'-?\s?'
                      r'\d{3}\D?'
                      r'-?\s?'
                      r'\d*'
                      r'-?\s?'
                      r'\D*\d*'
                      r'-?\s?'
                      r'\D*', line_or_drawing):
            # если количество "-" не больше трёх, то это номер сосуда
            if not line_or_drawing.count('-') > 3:
                line = re.findall(r'[AАBВCСDHНMМ]'
                                  r'\d{1,2}'
                                  r'-{1,2}?\s?'
                                  r'\d{3,4}'
                                  r'-?\s?'
                                  r'\D{2}'
                                  r'-?\s?'
                                  r'\d{3}\D?', line_or_drawing)
            else:
                line = re.findall(r'[AАBВCСDHНMМ]'
                                  r'\d{1,2}'
                                  r'-{1,2}?\s?'
                                  r'\d{3,4}'
                                  r'-?\s?'
                                  r'\D{2}'
                                  r'-?\s?'
                                  r'\d{3}\D?'
                                  r'-?\s?'
                                  r'\d*'
                                  r'-?\s?'
                                  r'\D*\d*'
                                  r'-?\s?'
                                  r'\D*', line_or_drawing)
            line = line[0].replace('\n', '')
            # вставляем на первое место номер линии
            for number_table in pure_data_table.keys():
                # в названия столбцов
                pure_data_table[number_table][0].insert(0, 'Line')
                # в данные
                for index, row in enumerate(pure_data_table[number_table][1]):
                    pure_data_table[number_table][1][index].insert(0, line)
    except UnboundLocalError:
        logger_with_user.error(f'В репорте {report_number} таблице {number_dirty_table} какая-то ошибка! А именно:\n'
                               f'{traceback.format_exc()}')
        stop_stop = True
        return stop_stop
    try:
        # если это чертёж
        if re.findall(r'[KК][EЕ]01-.+|[TТ]R01[-.].+', line_or_drawing):
            drawing = re.findall(r'[KК][EЕ]01-.+|[TТ]R01[-.].+', line_or_drawing)
            drawing = drawing[0].replace('\n', '')
            # вставляем на второе место номер чертежа
            for number_table in pure_data_table.keys():
                # в названия столбцов
                pure_data_table[number_table][0].insert(1, 'Drawing')
                # в данные
                for index, row in enumerate(pure_data_table[number_table][1]):
                    pure_data_table[number_table][1][index].insert(1, drawing)
        return pure_data_table
    except UnboundLocalError:
        logger_with_user.error(f'В репорте {report_number} таблице {number_dirty_table} какая-то ошибка! А именно:\n'
                               f'{traceback.format_exc()}')
        stop_stop = True
        return stop_stop


# удаляем Annex и Note в конце
def clean_up_end(pure_data_table: dict) -> dict:
    for number_table in pure_data_table.keys():
        index_row_for_delete = list()
        for index_row, row in enumerate(pure_data_table[number_table][1]):
            if len(set(row)) == 1:
                index_row_for_delete.append(index_row)
            for cell in row:
                if 'EXAMINED' in cell.upper():
                    end_row = len(pure_data_table[number_table][1])
                    for i in range(index_row, end_row):
                        index_row_for_delete.append(i)
                    break

        index_row_for_delete.reverse()
        # удаляем по индексу лишние, последние строки без данных
        for delete_index in index_row_for_delete:
            pure_data_table[number_table][1].pop(delete_index)
    return pure_data_table


# определение номера unit
def unit_definition(pure_data_table: dict, number_report: str) -> str:
    # print(pure_data_table)
    # список линий и чертежей
    line = list()
    drawing = list()
    for number_table in pure_data_table.keys():
        for index_name_column, name_column in enumerate(pure_data_table[number_table][0]):
            if name_column == 'Line':
                for row in pure_data_table[number_table][1]:
                    line.append(row[index_name_column])
            if name_column == 'Drawing':
                for row in pure_data_table[number_table][1]:
                    drawing.append(row[index_name_column])
    # убираем повторы найденных линий и чертежей
    set_line = list(set(line))
    set_drawing = list(set(drawing))
    # список определённых unit-ов
    unit_line = list()
    unit_drawing = list()
    # итоговый список найденных unit-ов
    unit = list()
    # определение, что номер линии - это сосуд
    vessel = False
    # выбираем номера из найденных линий и чертежей
    if set_line:
        for i in set_line:
            # если количество "-" не больше трёх, то это номер сосуда
            if not i.count('-') > 3:
                vessel = True
            first_hyphen = i.find('-')
            second_hyphen = i.find('-', i.find('-') + 1)
            if vessel:
                if first_hyphen < 0 or second_hyphen < 0:
                    unit_line.append(i)
                    logger_with_user.info(f'Проверь unit в {number_report}')
                else:
                    unit_line.append(i[first_hyphen + 1:second_hyphen])
            else:
                if first_hyphen < 0 or second_hyphen < 0:
                    unit_line.append(i)
                    logger_with_user.info(f'Проверь unit в {number_report}')
                else:
                    unit_line.append(i[first_hyphen + 1:second_hyphen - 1])
    if set_drawing:
        for i in set_drawing:
            second_hyphen = i.find('-', i.find('-') + 1)
            third_hyphen = i.find('-', i.find('-', i.find('-') + 1) + 1)
            unit_drawing.append(i[second_hyphen + 1:third_hyphen])
    if unit_line and unit_drawing:
        if list(set(unit_line) ^ set(unit_drawing)):
            for i in unit_line:
                unit.append(i)
            logger_with_user.info(f'Проверь записанный unit в {number_report}')
        else:
            unit = unit_drawing
    elif unit_line or unit_drawing:
        if unit_line:
            for i in unit_line:
                unit.append(i)
        elif unit_drawing:
            for i in unit_drawing:
                unit.append(i)
    else:
        # возвращаем пустой список
        unit = '-'
        logger_with_user.info(f'Проверь unit в {number_report}')
    # убираем повторы
    unit = list(set(unit))
    # избавляемся от ложных unit-ов - только трёхзначное число
    if len(unit) > 1:
        for i in unit:
            if len(i) > 2 and i.isnumeric():
                unit = i[:3]
    else:
        if not unit == ['-']:
            unit = str(unit[0])
        else:
            unit = '-'
    if type(unit) != str:
        logger_with_user.info(f'В {number_report} целый список unit - этого не может быть!')
        breakbreak = True
        return breakbreak
    return unit


# приводим дату репорта к нормальному виду
# 25.06-29.06.2021
# 30 June – 03 July, 2021
# 03-05 July, 2021
# 24 & 25 July, 2021
# 23, 26.08.2021
# 1-2.09.2021
# 3.09.2021
# 11-14 September, 2021
# 19 - 20 January 2021
def format_date_to_normal_form(date: str) -> str:
    date = date.replace(' ', '.')
    date = date.replace(',', '.')
    date = date.replace('.', '.')
    date = date.replace('/', '.')
    date = date.replace('-', '.')
    date = date.replace('_', '.')
    # разделяем дату на составляющие
    split_date = date.split('.')
    day = split_date[0]
    month = split_date[1]
    year = split_date[2]
    for index, day_month_year in enumerate(split_date):
        # убираем лишние окончания
        if 'th' in day_month_year:
            split_date[index] = split_date[index].replace('th', '')
        elif 'Th' in day_month_year:
            split_date[index] = split_date[index].replace('Th', '')
        elif 'TH' in day_month_year:
            split_date[index] = split_date[index].replace('TH', '')
        elif 'tH' in day_month_year:
            split_date[index] = split_date[index].replace('tH', '')
    # если в дате чёткое разделение на дату, месяц, год
    if len(split_date) == 3:
        # убираем диапазон дат
        if '-' in split_date[0]:
            ind = split_date[0].index('-')
            split_date[0] = split_date[0][ind + 1:]
        # оставляем последние дни
        day = split_date[0][-2:]
        # определяем месяц, если он словами
        for mon in month_alpha.keys():
            if month_alpha[mon] in split_date[1].upper():
                month = month_numeric[mon]
        # если месяц это числа
        if split_date[1].isdigit():
            if len(split_date[1]) == 2:
                month = split_date[1]
            elif len(split_date[1]) < 2:
                month = f'0{split_date[1]}'
            elif len(split_date[1]) > 2:
                if int(split_date[1][1:]) < 13:
                    month = split_date[1][1:]
                if int(split_date[1][1:]) > 12:
                    if int(split_date[1][:-1]) < 13:
                        month = split_date[1][:-1]
        # определяем год
        if len(split_date[2]) == 4:
            year = split_date[2]
        if len(split_date[2]) == 2:
            year = f'20{split_date[2]}'
    # если в дате переход из одного месяца в другой
    if len(split_date) == 4:
        # убираем диапазон дат
        if '-' in split_date[1]:
            ind = split_date[1].index('-')
            # оставляем последние дни
            day = split_date[1][ind + 1:]
        # определяем месяц, если он словами
        for mon in month_alpha.keys():
            if month_alpha[mon] in split_date[2].upper():
                month = month_numeric[mon]
        # если месяц это числа
        if split_date[2].isdigit():
            if len(split_date[2]) == 2:
                month = split_date[2]
            elif len(split_date[2]) < 2:
                month = f'0{split_date[1]}'
            elif len(split_date[2]) > 2:
                if int(split_date[2][1:]) < 13:
                    month = split_date[1][1:]
                if int(split_date[2][1:]) > 13:
                    if int(split_date[2][:-1]) < 13:
                        month = split_date[2][:-1]
        # определяем год
        if len(split_date[3]) == 4:
            year = split_date[3]
        if len(split_date[3]) == 2:
            year = f'20{split_date[3]}'
    # конечная дата
    true_date = f'{day}.{month}.{year}'
    return true_date
